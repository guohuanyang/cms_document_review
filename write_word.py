# -*- coding: utf-8 -*-
# Author:小涩
# Time: 2021/3/10 16:02
import json
import os
import re
import shutil
import sys
import time
from copy import deepcopy

import docx
import openpyxl
import pythoncom
import win32com.client as win32
from openpyxl.styles import Alignment, Border, Side
from config import dir_name

pythoncom.CoInitialize()


def open_doc(file_path):
    try:
        word = win32.gencache.EnsureDispatch("Word.Application")
    except Exception as err:
        print(err)
        MODULE_LIST = [m.__name__ for m in sys.modules.values()]
        for module in MODULE_LIST:
            # Remove cache and try again.
            MODULE_LIST = [m.__name__ for m in sys.modules.values()]
            for module in MODULE_LIST:
                if re.match(r'win32com\.gen_py\..+', module):
                    del sys.modules[module]
            shutil.rmtree(os.path.join(os.environ.get('LOCALAPPDATA'), 'Temp', 'gen_py'))
            word = win32.gencache.EnsureDispatch("Word.Application")
    word.Visible = False
    word.DisplayAlerts = 0
    doc = word.Documents.Open(file_path)
    return doc, word


def get_excel_selection(file_path, word, doc, len, company):
    excel = win32.DispatchEx("Excel.Application")
    book = excel.Workbooks.Open(file_path)
    selection = word.Selection
    myRange1 = doc.Range(0, 0)
    myRange1.InsertBefore(company)
    selection.EndKey()
    for i in range(1, len + 1):
        book.Sheets(str(i)).UsedRange.Copy()
        selection.PasteExcelTable(LinkedToExcel=False, WordFormatting=False, RTF=False)
        selection.InsertBreak(win32.constants.wdPageBreak)

    book.Close()
    selection.WholeStory()
    selection.Font.Size = 10


sheet = []


class WriteExcel:
    """
    按照模版写入excel文件
    """

    def __init__(self, header=None):
        wb = openpyxl.Workbook()
        self.wb = wb
        # self.ws = None
        self.ws = self.wb.active
        self.line_num = 1
        self.merge_level = 2
        self.header = header or list()
        self.start = 0

    def create_sheet(self, sheet_name):
        self.ws = self.wb.create_sheet(sheet_name)
        self.start = 0
        self.line_num = 1
        self.merge_level = 2

    def write_by_xy(self, data):
        end = data['level'] * self.merge_level
        if end > 1:
            self.start = end - 1
            self.ws.cell(self.line_num + 1, self.start + 1, data['ins_name'])
            self.ws.cell(self.line_num + 1, self.start + 1).alignment = Alignment(horizontal='center',
                                                                                  vertical='center')
            data['cell_1'] = {
                "row": self.line_num + 1,
                "col": self.start + 1,
            }
            self.ws.cell(self.line_num + 1, self.start + 2, "%s%%" % data['ratio'])
            self.ws.cell(self.line_num + 1, self.start + 2).alignment = Alignment(horizontal='center',
                                                                                  vertical='center')

            data['cell_2'] = {
                "row": self.line_num + 1,
                "col": self.start + 2,
            }
            # self.ws.cell(self.line_num + 1, start + 3, data['msg'] or '正常')
        else:
            self.ws.cell(self.line_num + 1, self.start + 1, data['ins_name'])
            self.ws.cell(self.line_num + 1, self.start + 1).alignment = Alignment(horizontal='center',
                                                                                  vertical='center')

            data['cell_1'] = {
                "row": self.line_num + 1,
                "col": self.start + 1,
            }

    def write_header(self):
        for index, level in enumerate(self.header, 1):
            if index == 1:
                self.ws.cell(1, index, level)
            else:
                col = (index - 1) * self.merge_level
                self.ws.cell(1, col, level)
                self.ws.merge_cells(start_row=1, end_row=1, start_column=col, end_column=col + 1)

    def dfs_write(self, root):
        if not root:
            return
        data = root['data']
        self.write_by_xy(data)
        if len(root['children']) == 0:
            self.line_num += 1
        for children in root['children']:
            self.dfs_write(children)
        self.merge_children(root)

    def merge_children(self, root):
        children = root['children']
        # 孩子节点数量大于1的
        if len(children) > 1:
            # 该节点的行
            start_row_num = root['data']['cell_1']['row']
            # 最后一个孩子节点的行
            end_row_num = children[-1]['data']['cell_1']['row']
            # 该节点的列
            start_col_num = root['data']['cell_1']['col']
            # 根节点
            if start_col_num == 1:
                self.ws.merge_cells(start_row=start_row_num, end_row=end_row_num, start_column=start_col_num,
                                    end_column=start_col_num)
                root['data']['cell_1'] = {
                    "row": end_row_num,
                    "col": start_col_num,
                }
            else:
                self.ws.merge_cells(start_row=start_row_num, end_row=end_row_num, start_column=start_col_num,
                                    end_column=start_col_num)
                root['data']['cell_1'] = {
                    "row": end_row_num,
                    "col": start_col_num,
                }
                self.ws.merge_cells(start_row=start_row_num, end_row=end_row_num, start_column=start_col_num + 1,
                                    end_column=start_col_num + 1)
                root['data']['cell_2'] = {
                    "row": end_row_num,
                    "col": start_col_num + 1,
                }
        elif len(children) == 1:
            start_row_num = root['data']['cell_1']['row']
            end_row_num = children[0]['data']['cell_1']['row']
            start_col_num = root['data']['cell_1']['col']
            if start_col_num == 1:
                self.ws.merge_cells(start_row=start_row_num, end_row=end_row_num, start_column=start_col_num,
                                    end_column=start_col_num)
                root['data']['cell_1'] = {
                    'row': end_row_num,
                    'col': start_col_num
                }
            else:
                self.ws.merge_cells(start_row=start_row_num, end_row=end_row_num, start_column=start_col_num,
                                    end_column=start_col_num)
                root['data']['cell_1'] = {
                    'row': end_row_num,
                    'col': start_col_num
                }
                self.ws.merge_cells(start_row=start_row_num, end_row=end_row_num, start_column=start_col_num + 1,
                                    end_column=start_col_num + 1)
                root['data']['cell_2'] = {
                    'row': end_row_num,
                    'col': start_col_num + 1
                }

    def save(self, company, save_dir=dir_name):
        file_name = "test" + '.xlsx'
        if not os.path.exists(dir_name):
            os.mkdir(dir_name)
        excel_file_path = "%s/%s" % (save_dir, file_name)
        self.wb.save(excel_file_path)
        self.wb.close()
        doc = docx.Document()
        # run = doc.add_heading('', level=1).add_run(company)
        doc.save("%s/result.docx" % save_dir)

        return os.path.join(os.getcwd(), "result/test.xlsx"), os.path.join(os.getcwd(), "result/result.docx"), company

    def add_title(self, level):
        if level == 0:
            self.ws.cell(1, 1, "名称")
        else:
            self.ws.cell(1, 1, "第%s层级" % level)
        self.ws.cell(1, 1).alignment = Alignment(horizontal='center', vertical='center')
        level += 1
        for i in range(2, 10, 2):
            self.ws.merge_cells(start_row=1, end_row=1, start_column=i,
                                end_column=i + 1)
            self.ws.cell(1, i, "第%s层级" % level)
            self.ws.cell(1, i).alignment = Alignment(horizontal='center', vertical='center')
            level += 1

    def changeHW(self):
        # width = 10.0
        # height = width * (2.2862 / 0.3612)

        # print("row:", self.ws.max_row, "column:", self.ws.max_column)
        for i in range(1, self.ws.max_row + 1):
            for j in range(1, self.ws.max_column + 1):
                # self.ws.column_dimensions[get_column_letter(i)].width = width
                # self.ws.cell(i, j).alignment = Alignment()

                self.ws.cell(row=i, column=j).alignment = Alignment(wrapText=True, horizontal='center',
                                                                    vertical='center')
        border_type = Side(border_style="thin", color="000000")
        border = Border(left=border_type,
                        right=border_type,
                        top=border_type,
                        bottom=border_type,
                        diagonal=border_type,
                        diagonal_direction=0,
                        outline=border_type,
                        vertical=border_type,
                        horizontal=border_type
                        )
        for row in self.ws['A1:I%s' % self.ws.max_row]:
            for cell in row:
                cell.border = border  # A5:D6区域单元格设置边框


def get_json_list(big_json, l, level):
    if level == 5:
        l.append(big_json)
    level -= 1
    for child in big_json["children"]:
        if level != 0:
            get_json_list(child, l, level)
        # else:
        #     get_json_list(child, l, 5)
    if level == 0:
        new_root = deepcopy(big_json)
        if new_root["children"]:
            get_json_list(new_root, l, 5)
        big_json["children"] = []


def update_level(data, level):
    data["data"]["level"] = level
    for child in data["children"]:
        update_level(child, level + 1)


def main(data):
    company = data.get("data").get("ins_name")
    company = "附件二：%s公司的穿透核查情况" % company
    all_jsons = []
    get_json_list(data, all_jsons, 5)
    obj1 = WriteExcel()
    for ind, _json in enumerate(all_jsons, 1):
        print(_json)
        level = _json["data"]["level"]
        update_level(_json, 0)
        obj1.create_sheet(str(ind))
        obj1.dfs_write(_json)
        obj1.add_title(level)
        obj1.changeHW()
        time.sleep(0.1)
        print(ind)
    excel_path, docx_path, company = obj1.save(company)
    doc, word = open_doc(docx_path)
    get_excel_selection(excel_path, word, doc, len(all_jsons), company)
    doc.Save()
    doc.Close()
    os.remove(excel_path)


if __name__ == '__main__':
    with open("./dfs.json", "r", encoding="utf8") as f:
        test_data = json.loads(f.read())
    main(test_data)
